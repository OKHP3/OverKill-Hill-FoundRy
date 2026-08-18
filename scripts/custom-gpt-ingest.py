#!/usr/bin/env python3
"""Create a source-preserving inventory and text extraction for Custom GPT research.

The default input and output locations are intentionally inside the Git-ignored
temporary ingestion workspace. This script does not modify its source files.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

try:
    import pdfplumber
except ImportError:  # The pypdf fallback remains usable.
    pdfplumber = None


SUPPORTED = {".docx", ".md", ".pdf", ".txt", ".xlsx", ".yaml", ".yml"}
DEFAULT_INPUT = Path("custom-gpts/ingestion")
DEFAULT_OUTPUT_NAME = "_processed"


@dataclass
class SourceRecord:
    source_id: str
    relative_path: str
    extension: str
    bytes: int
    sha256: str
    modified_utc: str
    units: int | None
    extraction_status: str
    extracted_characters: int
    extracted_text_sha256: str | None
    extractor: str
    notes: list[str]
    extracted_path: str | None


def content_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def source_id(relative_path: Path, digest: str) -> str:
    safe_name = re.sub(r"[^A-Za-z0-9]+", "-", relative_path.stem).strip("-").lower()
    return f"{safe_name[:48] or 'source'}-{digest[:10]}"


def markdown_front_matter(record: SourceRecord) -> str:
    return "\n".join(
        [
            "---",
            f"source_id: {record.source_id}",
            f"source_path: {json.dumps(record.relative_path)}",
            f"source_sha256: {record.sha256}",
            f"source_modified_utc: {record.modified_utc}",
            f"extraction_status: {record.extraction_status}",
            f"extractor: {record.extractor}",
            "---",
            "",
            "# Extracted source",
            "",
        ]
    )


def markdown_from_docx(path: Path) -> tuple[str, int, list[str]]:
    document = Document(path)
    blocks: list[str] = []
    notes: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        heading_match = re.search(r"heading\s+([1-6])", style_name)
        if heading_match:
            blocks.append(f"{'#' * int(heading_match.group(1))} {text}")
        else:
            blocks.append(text)
    for index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            blocks.append(f"\n## Table {index}\n")
            blocks.append(" | ".join(rows[0]))
            blocks.append(" | ".join("---" for _ in rows[0]))
            blocks.extend(" | ".join(row) for row in rows[1:])
    if document.inline_shapes:
        notes.append(f"contains {len(document.inline_shapes)} inline image(s) not transcribed as text")
    if not blocks:
        notes.append("no paragraph or table text extracted")
    return "\n\n".join(blocks).strip() + "\n", len(document.paragraphs), notes


def markdown_from_pdf(path: Path) -> tuple[str, int, str, list[str]]:
    notes: list[str] = []
    blocks: list[str] = []
    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as document:
                for index, page in enumerate(document.pages, start=1):
                    text = page.extract_text() or ""
                    blocks.append(f"## Page {index}\n\n{text.strip()}")
                return "\n\n".join(blocks).strip() + "\n", len(document.pages), "pdfplumber", notes
        except Exception as error:  # Fall through to pypdf and record why.
            notes.append(f"pdfplumber fallback: {type(error).__name__}: {error}")
    document = PdfReader(str(path))
    for index, page in enumerate(document.pages, start=1):
        blocks.append(f"## Page {index}\n\n{(page.extract_text() or '').strip()}")
    return "\n\n".join(blocks).strip() + "\n", len(document.pages), "pypdf", notes


def markdown_from_text(path: Path) -> tuple[str, int, list[str]]:
    for encoding in ("utf-8", "utf-8-sig", "cp1252"):
        try:
            text = path.read_text(encoding=encoding)
            return text.strip() + "\n", text.count("\n") + 1, []
        except UnicodeDecodeError:
            continue
    return "", 0, ["could not decode as UTF-8, UTF-8 BOM, or Windows-1252"]


def markdown_from_xlsx(path: Path) -> tuple[str, int, list[str]]:
    """Extract visible workbook values as reviewable Markdown tables."""
    workbook = load_workbook(path, read_only=True, data_only=False)
    blocks: list[str] = []
    row_count = 0
    notes: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            rows = [
                ["" if value is None else str(value).replace("|", "\\|").replace("\n", "<br>")
                 for value in row]
                for row in worksheet.iter_rows(values_only=True)
            ]
            while rows and not any(rows[-1]):
                rows.pop()
            if not rows:
                notes.append(f"worksheet {worksheet.title!r} contains no visible values")
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            row_count += len(rows)
            blocks.extend(
                [
                    f"## Worksheet: {worksheet.title}",
                    "",
                    " | ".join(rows[0]),
                    " | ".join("---" for _ in rows[0]),
                    *(" | ".join(row) for row in rows[1:]),
                    "",
                ]
            )
    finally:
        workbook.close()
    if not blocks:
        notes.append("no worksheet values extracted")
    return "\n".join(blocks).strip() + "\n", row_count, notes


def extract_source(path: Path, relative_path: Path, output_dir: Path) -> SourceRecord:
    extension = path.suffix.lower()
    digest = content_hash(path)
    source = SourceRecord(
        source_id=source_id(relative_path, digest),
        relative_path=relative_path.as_posix(),
        extension=extension,
        bytes=path.stat().st_size,
        sha256=digest,
        modified_utc=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
        units=None,
        extraction_status="failed",
        extracted_characters=0,
        extracted_text_sha256=None,
        extractor="",
        notes=[],
        extracted_path=None,
    )
    try:
        if extension == ".docx":
            text, source.units, source.notes = markdown_from_docx(path)
            source.extractor = "python-docx"
        elif extension == ".pdf":
            text, source.units, source.extractor, source.notes = markdown_from_pdf(path)
        elif extension == ".xlsx":
            text, source.units, source.notes = markdown_from_xlsx(path)
            source.extractor = "openpyxl"
        else:
            text, source.units, source.notes = markdown_from_text(path)
            source.extractor = "direct text read"
        source.extracted_characters = len(text.strip())
        source.extracted_text_sha256 = hashlib.sha256(
            re.sub(r"\s+", " ", text).strip().encode("utf-8")
        ).hexdigest()
        if source.extracted_characters < 40:
            source.notes.append("very little text extracted; inspect for scan, image-only content, or unsupported structure")
            source.extraction_status = "needs_review"
        else:
            source.extraction_status = "extracted"
        extraction_file = output_dir / "extracted" / f"{source.source_id}.md"
        extraction_file.parent.mkdir(parents=True, exist_ok=True)
        extraction_file.write_text(markdown_front_matter(source) + text, encoding="utf-8")
        source.extracted_path = extraction_file.relative_to(output_dir).as_posix()
    except Exception as error:
        source.notes.append(f"{type(error).__name__}: {error}")
        source.extractor = source.extractor or "not completed"
    return source


def write_inventory(records: list[SourceRecord], output_dir: Path) -> None:
    exact_duplicate_groups: dict[str, list[str]] = {}
    for record in records:
        if record.extracted_text_sha256:
            exact_duplicate_groups.setdefault(record.extracted_text_sha256, []).append(record.relative_path)
    inventory = {
        "generated_utc": datetime.now(tz=UTC).isoformat(),
        "source_count": len(records),
        "by_extension": {extension: sum(record.extension == extension for record in records) for extension in SUPPORTED},
        "records": [asdict(record) for record in records],
        "exact_content_duplicate_groups": [
            paths for paths in exact_duplicate_groups.values() if len(paths) > 1
        ],
    }
    (output_dir / "inventory.json").write_text(json.dumps(inventory, indent=2) + "\n", encoding="utf-8")
    lines = [
        "# Temporary source inventory",
        "",
        "This inventory is generated from the temporary ingestion workspace. It records extraction status and provenance, not an interpretation of the source material.",
        "",
        "| Source | Type | Units | Characters | Status | Notes |",
        "| --- | --- | ---: | ---: | --- | --- |",
    ]
    for record in records:
        notes = "; ".join(record.notes).replace("|", "\\|") or ""
        lines.append(
            f"| `{record.relative_path}` | `{record.extension}` | {record.units or 0} | {record.extracted_characters} | {record.extraction_status} | {notes} |"
        )
    (output_dir / "inventory.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=None)
    parser.add_argument("--clean", action="store_true", help="replace only the generated output directory")
    parser.add_argument(
        "--include-root-files",
        action="store_true",
        help="include README.md and .gitignore when they are source artifacts",
    )
    arguments = parser.parse_args()
    input_dir = arguments.input.resolve()
    output_dir = (arguments.output or input_dir / DEFAULT_OUTPUT_NAME).resolve()
    if not input_dir.is_dir():
        parser.error(f"input directory does not exist: {input_dir}")
    if output_dir == input_dir or input_dir not in output_dir.parents:
        parser.error("output directory must be a child of, not equal to, the ingestion input directory")
    if arguments.clean and output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    records = []
    for path in sorted(input_dir.rglob("*")):
        if not path.is_file() or output_dir in path.parents:
            continue
        if not arguments.include_root_files and path.parent == input_dir and path.name in {"README.md", ".gitignore"}:
            continue
        if path.suffix.lower() not in SUPPORTED:
            continue
        records.append(extract_source(path, path.relative_to(input_dir), output_dir))
    write_inventory(records, output_dir)
    failed = sum(record.extraction_status == "failed" for record in records)
    review = sum(record.extraction_status == "needs_review" for record in records)
    print(f"Processed {len(records)} source files: {failed} failed, {review} need review.")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
